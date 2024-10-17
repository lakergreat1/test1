import os
import json
from fastapi import FastAPI, Request, UploadFile, File, Form, HTTPException
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from openai import AsyncOpenAI
from reports import CrownBrief, GeneralOccurrence
from narrative_guidelines import general_guideline, crown_brief_guideline, general_occurrence_guideline
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
import tempfile
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
import shutil

app = FastAPI()

# CORS configuration
origins = [
    "http://localhost",
    "http://localhost:8000",
    "http://127.0.0.1",
    "http://127.0.0.1:8000",
    "https://pd-report-production.up.railway.app",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Mount static files
app.mount("/static", StaticFiles(directory="static"), name="static")

# Set up Jinja2 templates
templates = Jinja2Templates(directory="templates")

# Initialize OpenAI client
client = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))

OCCURRENCE_TYPES = [
    "Domestic Dispute", "Impaired Driving", "Property Crime", "Traffic Incident",
    "Drug Related", "Informational Report", "Public Disorder", "Violent Crime",
    "Fraud and Financial", "Missing Person", "Sexual Offense", "Other"
]

REPORT_TYPES = ["Crown Brief", "General Occurrence"]

@app.get("/")
async def root(request: Request):
    return templates.TemplateResponse("index.html", {"request": request, "occurrence_types": OCCURRENCE_TYPES, "report_types": REPORT_TYPES})

@app.post("/transcribe")
async def transcribe_audio(file: UploadFile = File(...)):
    temp_file_path = "temp_audio.mp3"
    try:
        print(f"Received audio file: {file.filename}")
        # Save the uploaded file temporarily
        with open(temp_file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        print("Transcribing audio...")
        # Transcribe the audio using Whisper API
        with open(temp_file_path, "rb") as audio_file:
            transcription = await client.audio.transcriptions.create(
                model="whisper-1",
                file=audio_file
            )
        
        print("Transcription completed successfully")
        return JSONResponse(content={"transcription": transcription.text})
    
    except Exception as e:
        print(f"Error during transcription: {str(e)}")
        raise HTTPException(status_code=500, detail="Transcription failed")
    
    finally:
        # Ensure the file is closed before attempting to remove it
        if 'audio_file' in locals():
            audio_file.close()
        
        # Try to remove the temporary file
        try:
            os.remove(temp_file_path)
            print(f"Temporary file {temp_file_path} removed")
        except Exception as e:
            print(f"Warning: Could not remove temporary file {temp_file_path}: {str(e)}")

@app.post("/generate_report")
async def generate_report(occurrence_type: str = Form(...), report_type: str = Form(...), transcription: str = Form(...)):
    try:
        print(f"Generating report for {occurrence_type} - {report_type}")
        
        if report_type == "Crown Brief":
            narrative_guideline = crown_brief_guideline
        elif report_type == "General Occurrence":
            narrative_guideline = general_occurrence_guideline
        else:
            narrative_guideline = general_guideline

        system_prompt = f"""Generate a {report_type} for a {occurrence_type} incident based on the following transcription. 
        Format the output as a structured report according to the provided schema. 
        If a field is not known, use 'information missing from transcript'. DO NOT imagine or make up any information. if any information is not in the transcript set that as "information missing from transcript".
        For the narrative section, follow these guidelines:
        {narrative_guideline}"""
        
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": transcription}
        ]

        report_schema = CrownBrief if report_type == "Crown Brief" else GeneralOccurrence

        print("Sending request to OpenAI API...")
        completion = await client.beta.chat.completions.parse(
            model="gpt-4o",
            messages=messages,
            response_format=report_schema
        )

        report_content = completion.choices[0].message.parsed
        print("Report generated successfully")

        # Save the report to a JSON file
        filename = f"report_{occurrence_type.lower().replace(' ', '_')}.json"
        with open(filename, "w") as f:
            json.dump(report_content.model_dump(), f, indent=2)
        print(f"Report saved to {filename}")

        return JSONResponse(content={"report": report_content.model_dump()})
    
    except Exception as e:
        print(f"Error generating report: {str(e)}")
        raise HTTPException(status_code=500, detail="Report generation failed")

@app.post("/edit_report")
async def edit_report(report: str = Form(...), instructions: str = Form(...), report_type: str = Form(...)):
    try:
        print("Editing report based on instructions")
        client = AsyncOpenAI()

        report_schema = CrownBrief if report_type == "Crown Brief" else GeneralOccurrence

        messages = [
            {"role": "system", "content": "You are a helpful assistant. Edit the given report according to the user's instructions. Ensure the edited report adheres to the provided schema."},
            {"role": "user", "content": f"Report:\n\n{report}\n\nInstructions:\n\n{instructions}"}
        ]

        completion = await client.beta.chat.completions.parse(
            model="gpt-4o-mini",
            messages=messages,
            response_format=report_schema
        )

        edited_report = completion.choices[0].message.parsed
        print("Report edited successfully")

        return JSONResponse(content={"edited_report": edited_report.model_dump()})
    
    except Exception as e:
        print(f"Error editing report: {str(e)}")
        raise HTTPException(status_code=500, detail="Report editing failed")

@app.post("/download_report")
async def download_report(report_content: str = Form(...), report_type: str = Form(...), format: str = Form(...)):
    try:
        if format == 'pdf':
            return await create_pdf_report(report_content, report_type)
        elif format == 'docx':
            return await create_word_report(report_content, report_type)
        else:
            raise HTTPException(status_code=400, detail="Invalid format specified")
    except Exception as e:
        print(f"Error generating {format} report: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate {format} report")

async def create_pdf_report(report_content: str, report_type: str):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                            rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=18)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Normal_LEFT',
                              parent=styles['Normal'],
                              alignment=0,
                              fontSize=10,
                              leading=14))

    story = []

    # Add title
    title = Paragraph(f"<b>{report_type} Report</b>", styles['Title'])
    story.append(title)
    story.append(Spacer(1, 0.25*inch))

    # Process report content
    lines = report_content.split('\n')
    for line in lines:
        if line.strip():  # Skip empty lines
            # Remove any unwanted characters (like black squares) that might appear at the end of lines
            clean_line = line.rstrip()
            p = Paragraph(clean_line, styles['Normal_LEFT'])
            story.append(p)
        else:
            story.append(Spacer(1, 0.1*inch))

    doc.build(story)
    buffer.seek(0)
    
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
        tmp.write(buffer.getvalue())
    
    return FileResponse(tmp.name, media_type="application/pdf", filename=f"{report_type.lower().replace(' ', '_')}_report.pdf")

async def create_word_report(report_content: str, report_type: str):
    doc = Document()
    
    # Add title
    title = doc.add_heading(f"{report_type} Report", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Set default font and size
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Process report content
    lines = report_content.split('\n')
    current_section = None
    
    for line in lines:
        line = line.strip()
        if line:
            if line.isupper() and line.endswith(':'):
                # This is a section header
                current_section = line
                p = doc.add_paragraph()
                p.add_run(line).bold = True
            elif ':' in line:
                # This is a field
                key, value = line.split(':', 1)
                p = doc.add_paragraph()
                p.add_run(f"{key.strip()}: ").bold = True
                p.add_run(value.strip())
            else:
                # This is probably part of the narrative
                doc.add_paragraph(line)
        elif current_section:
            # Add a small space between sections
            doc.add_paragraph()
            current_section = None
    
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        doc.save(tmp.name)
    
    return FileResponse(tmp.name, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=f"{report_type.lower().replace(' ', '_')}_report.docx")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="127.0.0.1", port=8000, reload=True)